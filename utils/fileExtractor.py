import io
import re
from fastapi import UploadFile, HTTPException

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 Megabytes

def extractTextFromDocBinary(content: bytes) -> str:
    """
    Best-effort ASCII/Unicode string extraction for legacy binary Word .doc files.
    """
    # Regex matching sequences of readable ASCII characters
    readable_segments = re.findall(br'[\x20-\x7E\s]{4,}', content)
    text = " ".join(seg.decode('ascii', errors='ignore') for seg in readable_segments)
    
    # Strip excessive whitespaces and header/footer metadata markers
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extractTextFromFile(file: UploadFile, content: bytes) -> str:
    """
    Validates file extension and size, then extracts and returns plain text.
    """
    filename = file.filename or ""
    lower_filename = filename.lower()
    
    # Size check
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=400, 
            detail=f"File size exceeds the maximum limit of 5MB. Got {(len(content)/(1024*1024)):.2f}MB."
        )
        
    # Extension validation and extraction
    if lower_filename.endswith(('.txt', '.md')):
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1')
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to decode text file: {str(e)}")
                
    elif lower_filename.endswith('.pdf'):
        try:
            from pypdf import PdfReader
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            extracted_text = "\n".join(text_parts).strip()
            if not extracted_text:
                raise HTTPException(status_code=400, detail="The PDF file appears to be empty or contains scanned images with no OCR text.")
            return extracted_text
        except Exception as e:
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")
            
    elif lower_filename.endswith('.docx'):
        try:
            import docx
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    
            # Extract table text as well
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                        
            extracted_text = "\n".join(text_parts).strip()
            if not extracted_text:
                raise HTTPException(status_code=400, detail="The DOCX file appears to be empty.")
            return extracted_text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")
            
    elif lower_filename.endswith('.doc'):
        # Legacy Word file best-effort handling
        try:
            extracted_text = extractTextFromDocBinary(content)
            if not extracted_text:
                raise HTTPException(status_code=400, detail="Failed to scrape readable text from legacy DOC file.")
            return extracted_text
        except Exception as e:
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(status_code=400, detail=f"Failed parsing legacy DOC: {str(e)}")
            
    else:
        raise HTTPException(
            status_code=400, 
            detail="Unsupported file extension. Only .txt, .md, .pdf, .docx, and .doc files are allowed."
        )
